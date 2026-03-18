"""
Compliance Report Generator
============================
Generates a compliance-ready, fully auditable Word document (.docx) describing
all data used within the Mergington High School Activities Management System.

Usage:
    python generate_compliance_report.py

Output:
    compliance_report.docx  (written to the repository root)
"""

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Apply a background colour to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_run(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.bold = True


def _add_table_with_headers(doc: Document, headers: list[str], rows: list[list[str]],
                             header_bg: str = "4472C4") -> None:
    """Add a formatted table with a coloured header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.text = header
        _set_cell_bg(cell, header_bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            row.cells[col_idx].text = str(value)

    doc.add_paragraph()  # spacing after table


# ---------------------------------------------------------------------------
# Report sections
# ---------------------------------------------------------------------------

def _cover_page(doc: Document, generated_at: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Mergington High School")
    run.bold = True
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Activities Management System")
    run2.bold = True
    run2.font.size = Pt(18)

    doc.add_paragraph()

    report_title = doc.add_paragraph()
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = report_title.add_run("Data Compliance & Audit Report")
    run3.bold = True
    run3.font.size = Pt(20)
    run3.font.color.rgb = RGBColor(0x17, 0x37, 0x5E)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {generated_at}\n")
    meta.add_run("Classification: Internal Use Only\n")
    meta.add_run("Version: 1.0")

    doc.add_page_break()


def _section_purpose(doc: Document) -> None:
    _add_heading(doc, "1. Purpose & Scope", 1)
    doc.add_paragraph(
        "This document provides a compliance-ready, fully auditable record of all "
        "data collected, stored, processed, and transmitted by the Mergington High "
        "School Activities Management System. It is intended to support internal "
        "audits, data protection reviews, and regulatory compliance assessments."
    )
    doc.add_paragraph(
        "Scope: All data elements present in the application as of the report "
        "generation date, including personally identifiable information (PII), "
        "system configuration data, and transient operational data."
    )


def _section_application_overview(doc: Document) -> None:
    _add_heading(doc, "2. Application Overview", 1)

    _add_heading(doc, "2.1 System Description", 2)
    doc.add_paragraph(
        "The Mergington High School Activities Management System is a web-based "
        "application that enables students to browse extracurricular activities and "
        "register their participation. The system consists of:"
    )
    items = [
        "A Python/FastAPI REST API backend serving activity data and handling registrations.",
        "A static HTML/CSS/JavaScript frontend served at the application root.",
        "An in-memory data store that is initialised at application start-up and reset on "
        "each server restart (no persistent database).",
    ]
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")

    _add_heading(doc, "2.2 Technology Stack", 2)
    stack_headers = ["Component", "Technology", "Version"]
    stack_rows = [
        ["Backend runtime", "Python", "3.13"],
        ["API framework", "FastAPI", "Latest (see requirements.txt)"],
        ["ASGI server", "Uvicorn", "Latest (see requirements.txt)"],
        ["Frontend", "HTML5 / CSS3 / Vanilla JavaScript", "—"],
        ["Data storage", "In-memory (Python dict)", "—"],
    ]
    _add_table_with_headers(doc, stack_headers, stack_rows)

    _add_heading(doc, "2.3 Deployment & Hosting", 2)
    doc.add_paragraph(
        "The application is designed for local development and educational use inside "
        "a GitHub Codespaces / Dev Container environment. It is not currently deployed "
        "to a production hosting environment. Port 8000 is exposed for local access."
    )


def _section_data_inventory(doc: Document) -> None:
    _add_heading(doc, "3. Data Inventory", 1)
    doc.add_paragraph(
        "The following table lists every distinct data element used by the application, "
        "its type, source, purpose, whether it constitutes PII, and the storage mechanism."
    )

    headers = [
        "Field Name",
        "Data Type",
        "Source",
        "Purpose",
        "PII?",
        "Storage",
        "Retention",
    ]
    rows = [
        [
            "Activity Name (key)",
            "String",
            "Static — hard-coded at startup",
            "Unique identifier for an extracurricular activity",
            "No",
            "In-memory dict key",
            "Duration of server process",
        ],
        [
            "description",
            "String",
            "Static — hard-coded at startup",
            "Human-readable summary of the activity",
            "No",
            "In-memory dict value",
            "Duration of server process",
        ],
        [
            "schedule",
            "String",
            "Static — hard-coded at startup",
            "Days and times the activity meets",
            "No",
            "In-memory dict value",
            "Duration of server process",
        ],
        [
            "max_participants",
            "Integer",
            "Static — hard-coded at startup",
            "Maximum enrolment cap for the activity",
            "No",
            "In-memory dict value",
            "Duration of server process",
        ],
        [
            "participants (email address)",
            "List of Strings (email)",
            "Student input via API POST /activities/{name}/signup",
            "Records which students are enrolled in an activity",
            "YES — school email address",
            "In-memory list within activity record",
            "Duration of server process; cleared on restart",
        ],
    ]
    _add_table_with_headers(doc, headers, rows)


def _section_pii(doc: Document) -> None:
    _add_heading(doc, "4. Personally Identifiable Information (PII)", 1)

    _add_heading(doc, "4.1 PII Elements Identified", 2)
    doc.add_paragraph(
        "The application collects one category of PII: student email addresses submitted "
        "when a student signs up for an activity. These are institutional (school-issued) "
        "email addresses in the format <username>@mergington.edu."
    )

    _add_heading(doc, "4.2 PII Risk Assessment", 2)
    risk_headers = ["Risk Factor", "Current Status", "Recommended Mitigation"]
    risk_rows = [
        [
            "No input validation on email field",
            "HIGH — any string accepted as email",
            "Implement RFC 5322 email format validation in the API",
        ],
        [
            "No duplicate-signup check",
            "MEDIUM — same email can be stored multiple times",
            "Check participants list before appending",
        ],
        [
            "max_participants not enforced",
            "MEDIUM — enrolment cap is not applied",
            "Return 400/409 if activity is at capacity",
        ],
        [
            "Data stored in plaintext in memory",
            "LOW (non-persistent; dev environment only)",
            "Use an encrypted database for production deployments",
        ],
        [
            "No authentication or authorisation",
            "HIGH — any caller can read all participants",
            "Implement authentication before production deployment",
        ],
        [
            "No audit logging",
            "MEDIUM — signup events are not logged",
            "Add structured logging of all data-mutation events",
        ],
    ]
    _add_table_with_headers(doc, risk_headers, risk_rows, header_bg="C00000")

    _add_heading(doc, "4.3 Data Minimisation", 2)
    doc.add_paragraph(
        "Only the student's email address is collected. No name, date of birth, address, "
        "or other personal data is requested. This aligns with data minimisation principles "
        "under GDPR Article 5(1)(c) and equivalent frameworks."
    )

    _add_heading(doc, "4.4 Lawful Basis for Processing", 2)
    doc.add_paragraph(
        "Processing of student email addresses is performed on the basis of the student's "
        "explicit consent (voluntary sign-up action). No data is shared with third parties."
    )


def _section_data_flows(doc: Document) -> None:
    _add_heading(doc, "5. Data Flows", 1)

    _add_heading(doc, "5.1 Data Flow Summary", 2)
    flow_headers = ["Flow ID", "From", "To", "Data Elements", "Protocol", "Encryption"]
    flow_rows = [
        [
            "DF-01",
            "Browser (student)",
            "API — GET /activities",
            "None (read-only request)",
            "HTTP",
            "None (dev environment)",
        ],
        [
            "DF-02",
            "API — GET /activities",
            "Browser (student)",
            "Activity names, descriptions, schedules, max_participants, participant count",
            "HTTP/JSON",
            "None (dev environment)",
        ],
        [
            "DF-03",
            "Browser (student)",
            "API — POST /activities/{name}/signup",
            "Email address (query parameter)",
            "HTTP",
            "None (dev environment)",
        ],
        [
            "DF-04",
            "API (signup handler)",
            "In-memory store",
            "Email address appended to participants list",
            "Internal (same process)",
            "N/A",
        ],
        [
            "DF-05",
            "API — POST /activities/{name}/signup",
            "Browser (student)",
            "Confirmation message string",
            "HTTP/JSON",
            "None (dev environment)",
        ],
    ]
    _add_table_with_headers(doc, flow_headers, flow_rows)

    _add_heading(doc, "5.2 External Data Transfers", 2)
    doc.add_paragraph(
        "No data is transferred to external systems, third-party services, or outside the "
        "local development environment. All data remains within the application process."
    )


def _section_api_endpoints(doc: Document) -> None:
    _add_heading(doc, "6. API Endpoint Catalogue", 1)
    doc.add_paragraph(
        "The following table documents every HTTP endpoint exposed by the application, "
        "the data it accepts, and the data it returns."
    )

    ep_headers = [
        "Method", "Path", "Input Data", "Output Data",
        "Authentication Required", "Notes",
    ]
    ep_rows = [
        [
            "GET",
            "/",
            "None",
            "HTTP 307 redirect to /static/index.html",
            "No",
            "Serves the frontend application",
        ],
        [
            "GET",
            "/activities",
            "None",
            "JSON object: all activities with description, schedule, "
            "max_participants, participants list",
            "No",
            "Exposes all participant email addresses — restrict in production",
        ],
        [
            "POST",
            "/activities/{activity_name}/signup",
            "Path param: activity_name (str)\nQuery param: email (str)",
            "JSON: {\"message\": \"Signed up <email> for <activity>\"}",
            "No",
            "No validation; no duplicate check; no capacity enforcement",
        ],
        [
            "GET",
            "/docs",
            "None",
            "Swagger UI (interactive API documentation)",
            "No",
            "Auto-generated by FastAPI; disable in production",
        ],
        [
            "GET",
            "/redoc",
            "None",
            "ReDoc API documentation",
            "No",
            "Auto-generated by FastAPI; disable in production",
        ],
    ]
    _add_table_with_headers(doc, ep_headers, ep_rows)


def _section_initial_data(doc: Document) -> None:
    _add_heading(doc, "7. Pre-loaded (Seed) Data", 1)
    doc.add_paragraph(
        "The application is initialised with the following hard-coded seed data at "
        "start-up. This data is defined in src/app.py and constitutes the complete "
        "initial state of the in-memory store."
    )

    seed_headers = [
        "Activity Name", "Description", "Schedule",
        "Max Participants", "Initial Participants",
    ]
    seed_rows = [
        [
            "Chess Club",
            "Learn strategies and compete in chess tournaments",
            "Fridays, 3:30 PM – 5:00 PM",
            "12",
            "michael@mergington.edu, daniel@mergington.edu",
        ],
        [
            "Programming Class",
            "Learn programming fundamentals and build software projects",
            "Tuesdays and Thursdays, 3:30 PM – 4:30 PM",
            "20",
            "emma@mergington.edu, sophia@mergington.edu",
        ],
        [
            "Gym Class",
            "Physical education and sports activities",
            "Mondays, Wednesdays, Fridays, 2:00 PM – 3:00 PM",
            "30",
            "john@mergington.edu, olivia@mergington.edu",
        ],
    ]
    _add_table_with_headers(doc, seed_headers, seed_rows)

    doc.add_paragraph(
        "NOTE: The email addresses in the seed data are fictional placeholder addresses "
        "used solely for demonstration purposes and do not represent real individuals."
    )


def _section_retention(doc: Document) -> None:
    _add_heading(doc, "8. Data Retention & Deletion", 1)

    _add_heading(doc, "8.1 Current Retention Policy", 2)
    doc.add_paragraph(
        "All data is held exclusively in the application process memory. There is no "
        "persistent storage (no database, no file system writes, no external service). "
        "All data — including any student email addresses collected during a session — "
        "is permanently and automatically deleted when the server process terminates."
    )

    _add_heading(doc, "8.2 Right to Erasure", 2)
    doc.add_paragraph(
        "Because data does not persist beyond the server process, a student's data is "
        "automatically erased on server restart. No manual deletion procedure is required "
        "in the current implementation. A production deployment would need to implement "
        "an explicit deletion endpoint to satisfy GDPR Article 17 (Right to Erasure)."
    )


def _section_access_controls(doc: Document) -> None:
    _add_heading(doc, "9. Access Controls", 1)

    doc.add_paragraph(
        "The application currently has no authentication or authorisation mechanisms. "
        "Any network-accessible client can read all activity data (including participant "
        "email addresses) and submit new sign-ups. The following controls are recommended "
        "before any production deployment:"
    )
    controls = [
        "Require authentication (e.g. OAuth 2.0 / OpenID Connect with school SSO) for "
        "all API endpoints.",
        "Restrict the GET /activities endpoint to return participant email addresses only "
        "to authorised staff roles.",
        "Rate-limit the POST signup endpoint to prevent abuse.",
        "Disable the /docs and /redoc Swagger endpoints in production.",
        "Enforce HTTPS for all data in transit.",
    ]
    for ctrl in controls:
        doc.add_paragraph(ctrl, style="List Bullet")


def _section_audit_log(doc: Document, generated_at: str) -> None:
    _add_heading(doc, "10. Audit Trail", 1)

    _add_heading(doc, "10.1 Report Generation Record", 2)
    audit_headers = ["Field", "Value"]
    audit_rows = [
        ["Report Generated", generated_at],
        ["Generated By", "generate_compliance_report.py (automated)"],
        ["Application Name", "Mergington High School Activities Management System"],
        ["Source File Audited", "src/app.py"],
        ["Report Version", "1.0"],
        ["Classification", "Internal Use Only"],
    ]
    _add_table_with_headers(doc, audit_headers, audit_rows, header_bg="375623")

    _add_heading(doc, "10.2 Application Audit Logging (Current State)", 2)
    doc.add_paragraph(
        "The application does not currently implement structured audit logging. "
        "All events (sign-ups, errors) are surfaced only via Uvicorn's default "
        "access log output to stdout/stderr. For compliance purposes, a production "
        "deployment must implement structured, tamper-evident audit logging of all "
        "data-mutation events."
    )


def _section_recommendations(doc: Document) -> None:
    _add_heading(doc, "11. Compliance Recommendations", 1)

    rec_headers = ["Priority", "Recommendation", "Relevant Standard"]
    rec_rows = [
        [
            "Critical",
            "Add authentication & authorisation before any production deployment",
            "GDPR Art. 32; ISO 27001 A.9",
        ],
        [
            "Critical",
            "Restrict participant data visibility to authorised roles only",
            "GDPR Art. 5(1)(f); FERPA",
        ],
        [
            "High",
            "Validate email format on signup (RFC 5322)",
            "Data quality / integrity controls",
        ],
        [
            "High",
            "Implement duplicate-signup prevention",
            "Data accuracy — GDPR Art. 5(1)(d)",
        ],
        [
            "High",
            "Enforce max_participants capacity limit",
            "Business logic integrity",
        ],
        [
            "High",
            "Enable HTTPS / TLS for all traffic",
            "GDPR Art. 32; NIST SP 800-53 SC-8",
        ],
        [
            "Medium",
            "Implement structured audit logging for all data-mutation events",
            "ISO 27001 A.12.4; SOX",
        ],
        [
            "Medium",
            "Migrate from in-memory storage to a persistent encrypted database for production",
            "GDPR Art. 32; ISO 27001 A.10",
        ],
        [
            "Medium",
            "Disable /docs and /redoc in production builds",
            "OWASP API Security Top 10 — API9",
        ],
        [
            "Low",
            "Add a data retention and deletion API endpoint",
            "GDPR Art. 17 (Right to Erasure)",
        ],
        [
            "Low",
            "Document and publish a Privacy Notice for students",
            "GDPR Art. 13; FERPA",
        ],
    ]
    _add_table_with_headers(doc, rec_headers, rec_rows, header_bg="833C00")


def _section_sign_off(doc: Document) -> None:
    _add_heading(doc, "12. Sign-Off", 1)
    doc.add_paragraph(
        "This report must be reviewed and approved by the designated Data Protection "
        "Officer (DPO) or equivalent role before being filed as an official compliance "
        "record."
    )

    sign_headers = ["Role", "Name", "Signature", "Date"]
    sign_rows = [
        ["Data Protection Officer", "", "", ""],
        ["System Owner", "", "", ""],
        ["IT / Security Lead", "", "", ""],
    ]
    _add_table_with_headers(doc, sign_headers, sign_rows, header_bg="595959")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def generate_report(output_path: Path) -> None:
    """Generate the compliance Word document and save it to *output_path*."""
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    doc = Document()

    # Document properties
    core_props = doc.core_properties
    core_props.title = "Data Compliance & Audit Report — Mergington High School Activities System"
    core_props.author = "Compliance Report Generator (automated)"
    core_props.subject = "Data Compliance"
    core_props.keywords = "compliance, audit, data, PII, GDPR"

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _cover_page(doc, generated_at)
    _section_purpose(doc)
    _section_application_overview(doc)
    _section_data_inventory(doc)
    _section_pii(doc)
    _section_data_flows(doc)
    _section_api_endpoints(doc)
    _section_initial_data(doc)
    _section_retention(doc)
    _section_access_controls(doc)
    _section_audit_log(doc, generated_at)
    _section_recommendations(doc)
    _section_sign_off(doc)

    doc.save(str(output_path))
    print(f"Compliance report written to: {output_path}")


if __name__ == "__main__":
    repo_root = Path(__file__).resolve().parent.parent
    output_file = repo_root / "compliance_report.docx"
    generate_report(output_file)
